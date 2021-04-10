import pickle
from os import listdir 
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm

from docx.oxml.ns import qn
import docx
#打开文档
from urllib3.connectionpool import xrange


path_model = 'Reactant_reactions.docx'
# document = docx.Document(path_model)
document = Document()
#增加表格
table = document.add_table(rows=1, cols=6,style='Table Grid')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Reactant_A'
hdr_cells[1].text = 'Reactant_B'
hdr_cells[2].text = 'Product'
hdr_cells[3].text = 'gibbs'
hdr_cells[4].text = 'homolumo'
hdr_cells[5].text = 'all_positive'

#pictures = [x for x in listdir("/root/qzh/e2e_reaction_test/homolumo") if x.endswith('.jpg')
#row_cells = table.add_row().cells
    
with open('Reactant_reactions_all.txt', 'rb') as fp:  
    b = pickle.load(fp) 

for i in range(len(b)):
    row_cells = table.add_row().cells
    for j in range(3):
        run=table.cell(i+1,j).paragraphs[0].add_run() 
        picture =run.add_picture("/root/qzh/e2e_reaction_test/homolumo/picture/"+str(b[i][j])+".png")
        picture.height=Cm(3)
        picture.width=Cm(3)
        
    row_cells[3].text = str(b[i][3])+'('+str(b[i][4])+')'
    row_cells[4].text = str(b[i][5])
    row_cells[5].text = str(b[i][6])

document.save(path_model)

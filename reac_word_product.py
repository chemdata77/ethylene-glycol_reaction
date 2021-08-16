import pickle
import cv2
from os import listdir 
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm

from docx.oxml.ns import qn
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH  
from docx.oxml.ns import qn  
from PIL import Image

with open('Product_reactions_gibbs_test_63.txt', 'rb') as fp: 
    q = pickle.load(fp) 
  
with open('Product_reactions_hl_test_63.txt', 'rb') as fp: 
    hl = pickle.load(fp)

with open('Product_reactions_all_test_81.txt', 'rb') as fp:
    al = pickle.load(fp)    
document = Document()

document.styles['Normal'].font.name = u'宋体'

document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

p1 = document.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = p1.add_run('乙二醇作为产物的反应表')
run1.font.name = '宋体'
run1.element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')

p2 = document.add_paragraph()
run2 = p2.add_run('说明：\n')
run2.font.name = '宋体'
run2.element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')

p3 = document.add_paragraph()
run3 = p2.add_run('\tGibbs自由能判据: ‘1’代表从热力学角度利于反应正向进行；‘-1’代表从热力学角度不利于反应正向进行。\n'+
            '\t前线轨道理论判据: ‘1’代表基于前线轨道理论利于反应正向进行；‘-1’代表基于前线轨道理论不利于反应正向进行。\n'+
            '\t是否同时满足: ‘1’代表同时满足Gibbs和前线轨道理论判据，利于反应正向进行；‘-1’代表不同时满足上述两判据。\n'+
             '\t数据来源：所有原始属性数据计算方法均为 B3LYP/6-31G(2df,p)，单位均为eV。')
run3.font.name = '宋体'
run3.element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')

#增加表格
table = document.add_table(rows=1, cols=6,style='Table Grid')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '反应物A'
hdr_cells[1].text = '反应物B'
hdr_cells[2].text = '产物C'
hdr_cells[3].text = 'Gibbs自由能判据(ΔG)'
hdr_cells[4].text = '前线轨道理论判据' 
hdr_cells[5].text = '是否同时满足'

path_model = '乙二醇作为产物的反应表.docx'


for j in range(len(q)): 
    row_cells = table.add_row().cells
    for k in range(3):
        run=table.cell(j+1,k).paragraphs[0].add_run() 
        picture =run.add_picture("/data/tp/picture_out/"+str(q[j][k])+".png")
             
        picture.height=Cm(2)
        picture.width=Cm(2)
        
        run4=table.cell(j+1,k).paragraphs[0].add_run('\nG:\n'+str(('%.5f' % q[j][4][k*2]))+'\nHOMO:\n'+str(('%.5f' % hl[j][4][k*2]))+'\nLOMO:\n'+str(('%.5f' % hl[j][4][k*2+1]))+'\n'+'factor:'+str(('%d' % q[j][4][k*2+1]))) 
    
    row_cells[3].text = str(q[j][3])+'\n'+'('+str(('%.5f' % q[j][4][6]))+')'

    row_cells[4].text = str(hl[j][3])
    row_cells[5].text = str(al[j][3])
document.save(path_model)
